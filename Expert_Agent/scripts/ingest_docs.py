
import os
import sys
import re
from typing import List, Dict, Any
from sqlalchemy.orm import Session
from sqlalchemy import text
from sentence_transformers import SentenceTransformer
import pypdf
import docx

# Add project root to path
sys.path.append(os.getcwd())

from data.database import SessionLocal, engine
from api.models import Document, DocumentChunk, Domain

def extract_text(filepath: str) -> str:
    """Extract text from PDF, DOCX, or MD files"""
    filename = os.path.basename(filepath).lower()
    try:
        if filename.endswith('.pdf'):
            reader = pypdf.PdfReader(filepath)
            return "\n".join([page.extract_text() for page in reader.pages])
        elif filename.endswith('.docx'):
            doc = docx.Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs])
        elif filename.endswith(('.md', '.txt')):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        print(f"  ⚠️ Error extracting {filepath}: {e}")
    return ""

def chunk_text(text: str, chunk_size: int = 600, overlap: int = 100) -> List[str]:
    """
    Paragraph-aware chunking with overlap.

    Handles edge cases:
    - Empty / whitespace-only text  -> returns []
    - Text shorter than chunk_size  -> returns as a single chunk
    - Very long paragraphs          -> split by sentence / newline boundaries
    - Lists / bullets               -> each item kept as its own segment
    """
    if not text or not text.strip():
        return []

    # Step 1: paragraph split first to preserve list/bullet structure
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    if not paragraphs:
        return []

    # Step 2: segment each paragraph into sentences / bullet lines
    segments: List[str] = []
    for para in paragraphs:
        para_clean = re.sub(r'[ \t]+', ' ', para)  # horizontal whitespace only
        parts = re.split(r'(?<=[.!?]) +|\n', para_clean)
        segments.extend(p.strip() for p in parts if p.strip())

    # Fast path: entire text fits in one chunk (short documents like test.txt)
    full_text = ' '.join(segments)
    if len(full_text) <= chunk_size:
        return [full_text]

    # Step 3: group into chunks with overlap for context continuity
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0

    for seg in segments:
        if current_len + len(seg) > chunk_size and current:
            chunk_str = ' '.join(current)
            chunks.append(chunk_str)
            # Carry last `overlap` chars into the next chunk for context
            carry = chunk_str[-overlap:] if len(chunk_str) > overlap else chunk_str
            current = [carry, seg]
            current_len = len(carry) + len(seg)
        else:
            current.append(seg)
            current_len += len(seg)

    if current:
        chunks.append(' '.join(current))

    return chunks

def run_ingestion():
    print("🚀 Starting Document Ingestion & Embedding Generation...")
    
    # 1. Initialize Model
    print("📥 Loading Embedding Model (all-MiniLM-L6-v2)...")
    model = SentenceTransformer('all-MiniLM-L6-v2')
    
    db: Session = SessionLocal()
    try:
        # 2. Get or infer the target domain
        domain = db.query(Domain).first()  # Use the first available domain
        if not domain:
            print("❌ No domain found in the database. Please run the seed scripts first.")
            return

        domain_id = domain.id
        print(f"🎯 Target Domain: {domain.name} (ID: {domain_id})")

        # 3. Scan docs directory
        # We look in the project root /docs (where the Sanofi file is)
        docs_dir = os.path.abspath(os.path.join(os.getcwd(), '..', 'docs'))
        if not os.path.exists(docs_dir):
            # Fallback to current Expert_Agent/docs if root/docs missing
            docs_dir = os.path.abspath(os.path.join(os.getcwd(), 'docs'))
            
        print(f"📂 Scanning directory: {docs_dir}")
        
        files = [f for f in os.listdir(docs_dir) if f.lower().endswith(('.md', '.pdf', '.docx', '.txt'))]
        print(f"📄 Found {len(files)} files to process.")

        for filename in files:
            filepath = os.path.join(docs_dir, filename)
            print(f"📖 Processing: {filename}...")
            
            content = extract_text(filepath)
            if not content:
                print(f"  ⚠️ No content extracted from {filename}, skipping.")
                continue
                
            # 4. Create/Update Document
            doc = db.query(Document).filter_by(domain_id=domain_id, source=filename).first()
            if not doc:
                doc = Document(
                    domain_id=domain_id,
                    title=filename.replace('_', ' ').replace('.pdf', '').replace('.docx', '').title(),
                    source=filename,
                    scope="public"
                )
                db.add(doc)
                db.flush() # Get doc.id
            else:
                # Update existing doc to ensure domain mapping
                doc.domain_id = domain_id
                # Remove old chunks to re-index
                db.query(DocumentChunk).filter_by(document_id=doc.id).delete()
            
            # 5. Chunk and Embed
            chunks = chunk_text(content)
            print(f"  🧩 Split into {len(chunks)} chunks.")
            
            for i, chunk_text_content in enumerate(chunks):
                # Generate embedding
                embedding = model.encode(chunk_text_content).tolist()
                
                new_chunk = DocumentChunk(
                    document_id=doc.id,
                    domain_id=domain_id,
                    content=chunk_text_content,
                    chunk_index=i,
                    embedding=embedding
                )
                db.add(new_chunk)
            
            db.commit()
            print(f"  ✅ {filename} indexed successfully.")

        print("🎉 Ingestion and Embedding process complete!")
        
    except Exception as e:
        print(f"❌ Error during ingestion: {e}")
        db.rollback()
        import traceback
        traceback.print_exc()
        
    finally:
        db.close()
        
    print("🚀 Running Knowledge Graph Induction Build...")
    from data.ontology_builder import OntologyBuilder
    ontology_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'knowledge_base', 'ontology.ttl'))
    if not os.path.exists(os.path.dirname(ontology_path)):
        ontology_path = os.path.abspath(os.path.join(os.getcwd(), 'knowledge_base', 'ontology.ttl'))
        os.makedirs(os.path.dirname(ontology_path), exist_ok=True)
    builder = OntologyBuilder(ontology_path=ontology_path)
    builder.build_from_directory(docs_dir)

if __name__ == "__main__":
    run_ingestion()
